"""
Open-source, rule-based chunker using LangChain's RecursiveCharacterTextSplitter.

Why this choice:
- Deterministic, no ML model download needed.
- Installs via pip alone on Windows.
- Handles mixed agricultural text (narrative + bullet lists + headings) well.
- Swap to SemanticChunker later by replacing this module's `chunk_text`.
"""
from __future__ import annotations

from pathlib import Path
from typing import List

from django.conf import settings
from langchain_text_splitters import RecursiveCharacterTextSplitter


def _splitter() -> RecursiveCharacterTextSplitter:
    return RecursiveCharacterTextSplitter(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
        length_function=len,
    )


def chunk_text(text: str) -> List[str]:
    text = (text or "").strip()
    if not text:
        return []
    return [c.strip() for c in _splitter().split_text(text) if c.strip()]


def extract_pdf_text(path: Path | str) -> str:
    """Extract text from a PDF file with pypdf. Returns empty string on failure."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        pages = []
        for page in reader.pages:
            try:
                pages.append(page.extract_text() or "")
            except Exception:
                continue
        return "\n\n".join(pages).strip()
    except Exception:
        return ""


def extract_docx_text(path: Path | str) -> str:
    """Extract text from a .docx Word file. Pulls paragraph text + table cell text."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(path))
        parts: list[str] = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts).strip()
    except Exception:
        return ""


def extract_xlsx_text(path: Path | str) -> str:
    """Extract text from a .xlsx Excel file (all sheets, all non-empty rows)."""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(str(path), read_only=True, data_only=True)
        parts: list[str] = []
        for sheet in wb.worksheets:
            parts.append(f"### Sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts).strip()
    except Exception:
        return ""


def extract_csv_text(path: Path | str) -> str:
    """Extract text from a .csv file (all rows, all columns)."""
    try:
        import csv as _csv
        parts: list[str] = []
        # Best-effort encoding handling
        with open(str(path), "r", encoding="utf-8", errors="replace", newline="") as fh:
            reader = _csv.reader(fh)
            for row in reader:
                cells = [c.strip() for c in row if c and c.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts).strip()
    except Exception:
        return ""


def extract_txt_text(path: Path | str) -> str:
    """Read a plain .txt file."""
    try:
        with open(str(path), "r", encoding="utf-8", errors="replace") as fh:
            return fh.read().strip()
    except Exception:
        return ""


# Map file extensions to extractors. To add a new format, drop a function above
# and add a row here — `document_text()` picks it up automatically.
_EXTRACTORS = {
    ".pdf":  extract_pdf_text,
    ".docx": extract_docx_text,
    ".xlsx": extract_xlsx_text,
    ".csv":  extract_csv_text,
    ".txt":  extract_txt_text,
}


def document_text(document) -> str:
    """
    Return the best available text body for a Document:
      1. Uploaded file (PDF / DOCX / XLSX / CSV / TXT — extracted)
      2. Fallback body built from metadata + keywords
    """
    if document.file and document.file.path:
        path = Path(document.file.path)
        if path.exists():
            extractor = _EXTRACTORS.get(path.suffix.lower())
            if extractor is not None:
                txt = extractor(path)
                if txt:
                    return txt
    # fallback body so every tagged CSV row still produces vectors
    return document.text_for_embedding()
